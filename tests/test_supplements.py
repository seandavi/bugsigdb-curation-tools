"""Unit tests for `bugsigdb_curation.supplements` -- fully offline.

`parse_supplement_refs`/`unpack_supplement_zip`/`supplement_to_text`/
`supplement_to_model_document` are pure (inline XML fixtures / in-memory
zip+xlsx+docx built inline, no network). `fetch_supplement_zip`/
`fetch_supplements` are covered with `pytest_httpx` mocks -- no live requests.
"""

from __future__ import annotations

import asyncio
import csv
import io
import zipfile

import docx
import httpx
import openpyxl
import pytest
from pytest_httpx import HTTPXMock

from bugsigdb_curation.supplements import (
    EUROPEPMC_SUPPLEMENTARY_FILES_URL,
    SupplementFile,
    SupplementRef,
    fetch_supplement_zip,
    fetch_supplements,
    parse_supplement_refs,
    supplement_to_model_document,
    supplement_to_text,
    unpack_supplement_zip,
)

XML_FIXTURE = """<?xml version="1.0" encoding="UTF-8"?>
<article xmlns:xlink="http://www.w3.org/1999/xlink">
  <body>
    <supplementary-material id="MOESM1">
      <label>Supplementary Material 1</label>
      <caption><p>Uncropped <italic>gel</italic> images.</p></caption>
      <media xlink:href="41598_2021_99379_MOESM1_ESM.pdf" mimetype="application" mime-subtype="pdf"/>
    </supplementary-material>
    <supplementary-material id="MOESM2">
      <label>Supplementary Table S1</label>
      <caption><p>Sample metadata.</p></caption>
      <media xlink:href="41598_2021_99379_MOESM2_ESM.xlsx"/>
    </supplementary-material>
  </body>
</article>
"""

NO_SUPPLEMENTS_XML = "<article><body><p>nothing here</p></body></article>"


# --- parse_supplement_refs ----------------------------------------------------------------


def test_parse_supplement_refs_extracts_id_filename_label_caption():
    refs = parse_supplement_refs(XML_FIXTURE)

    assert refs == [
        SupplementRef(
            supplement_id="MOESM1",
            filename="41598_2021_99379_MOESM1_ESM.pdf",
            label="Supplementary Material 1",
            caption="Uncropped gel images.",
        ),
        SupplementRef(
            supplement_id="MOESM2",
            filename="41598_2021_99379_MOESM2_ESM.xlsx",
            label="Supplementary Table S1",
            caption="Sample metadata.",
        ),
    ]


def test_parse_supplement_refs_returns_empty_list_when_none_present():
    assert parse_supplement_refs(NO_SUPPLEMENTS_XML) == []


# Real EuropePMC shape: caption is nested INSIDE <media> (not a direct child
# of <supplementary-material>), carries the <?suppdata-*?> processing
# instructions, and uses inline markup like <bold>. Verified against live
# PMC8497572 / PMC10590023 fullTextXML. A direct-child `find("caption")`
# returned "" on every real document (ledger: synthetic fixture missed it).
REAL_SHAPE_XML = """<?xml version="1.0" encoding="UTF-8"?>
<article xmlns:xlink="http://www.w3.org/1999/xlink">
  <body>
    <supplementary-material content-type="local-data" id="MOESM15" position="float">
      <media xmlns:xlink="http://www.w3.org/1999/xlink" xlink:href="40168_2023_1671_MOESM15_ESM.xlsx" position="float">
        <?suppdata-name 40168_2023_1671_MOESM15_ESM.xlsx?><?suppdata-size 69304?>
        <caption><p><bold>Additional file 15:</bold> Differential abundance results.</p></caption>
      </media>
    </supplementary-material>
  </body>
</article>
"""


def test_parse_supplement_refs_reads_caption_nested_in_media():
    refs = parse_supplement_refs(REAL_SHAPE_XML)
    assert len(refs) == 1
    ref = refs[0]
    assert ref.supplement_id == "MOESM15"
    assert ref.filename == "40168_2023_1671_MOESM15_ESM.xlsx"
    # inline <bold> flattened; caption recovered despite living under <media>
    assert ref.caption == "Additional file 15: Differential abundance results."


def test_parse_supplement_refs_falls_back_to_synthesized_id_and_none_filename():
    xml = "<article><body><supplementary-material><caption><p>no id, no media</p></caption></supplementary-material></body></article>"
    refs = parse_supplement_refs(xml)
    assert len(refs) == 1
    assert refs[0].supplement_id == "supp-0"
    assert refs[0].filename is None
    assert refs[0].label == ""
    assert refs[0].caption == "no id, no media"


# --- unpack_supplement_zip ----------------------------------------------------------------


def _build_zip(entries: dict[str, bytes], *, dirs: tuple[str, ...] = ()) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        for name in dirs:
            zf.writestr(zipfile.ZipInfo(name if name.endswith("/") else name + "/"), "")
        for name, content in entries.items():
            zf.writestr(name, content)
    return buf.getvalue()


def test_unpack_supplement_zip_skips_directory_entries_and_maps_media_type():
    zip_bytes = _build_zip(
        {"a.pdf": b"%PDF-1.4 fake", "b.csv": b"x,y\n1,2\n", "sub/c.png": b"fake-png-bytes"},
        dirs=("sub",),
    )
    files = unpack_supplement_zip(zip_bytes)

    by_name = {f.filename: f for f in files}
    assert set(by_name) == {"a.pdf", "b.csv", "c.png"}
    assert by_name["a.pdf"].media_type == "pdf"
    assert by_name["b.csv"].media_type == "csv"
    assert by_name["c.png"].media_type == "image"
    assert by_name["a.pdf"].raw_bytes == b"%PDF-1.4 fake"


@pytest.mark.parametrize(
    "filename,expected",
    [
        ("f.pdf", "pdf"),
        ("f.xlsx", "xlsx"),
        ("f.xls", "xlsx"),
        ("f.csv", "csv"),
        ("f.tsv", "tsv"),
        ("f.docx", "docx"),
        ("f.doc", "docx"),
        ("f.jpg", "image"),
        ("f.png", "image"),
        ("f.zip", "other"),
        ("f", "other"),
    ],
)
def test_unpack_supplement_zip_media_type_by_extension(filename, expected):
    zip_bytes = _build_zip({filename: b"data"})
    files = unpack_supplement_zip(zip_bytes)
    assert files[0].media_type == expected


# --- fetch_supplement_zip / fetch_supplements (network, mocked) --------------------------


def test_fetch_supplement_zip_returns_bytes_on_200(httpx_mock: HTTPXMock):
    zip_bytes = _build_zip({"a.pdf": b"data"})
    httpx_mock.add_response(
        url=EUROPEPMC_SUPPLEMENTARY_FILES_URL.format(pmcid="PMC8497572"),
        content=zip_bytes,
        headers={"Content-Type": "application/zip"},
    )

    async def run() -> bytes | None:
        async with httpx.AsyncClient() as client:
            return await fetch_supplement_zip("PMC8497572", client=client)

    assert asyncio.run(run()) == zip_bytes


def test_fetch_supplement_zip_returns_none_on_404(httpx_mock: HTTPXMock):
    httpx_mock.add_response(
        url=EUROPEPMC_SUPPLEMENTARY_FILES_URL.format(pmcid="PMC0000000"), status_code=404
    )

    async def run() -> bytes | None:
        async with httpx.AsyncClient() as client:
            return await fetch_supplement_zip("PMC0000000", client=client)

    assert asyncio.run(run()) is None


def test_fetch_supplement_zip_returns_none_on_non_zip_content_type(httpx_mock: HTTPXMock):
    httpx_mock.add_response(
        url=EUROPEPMC_SUPPLEMENTARY_FILES_URL.format(pmcid="PMC1111111"),
        content=b"<html>not a zip</html>",
        headers={"Content-Type": "text/html"},
    )

    async def run() -> bytes | None:
        async with httpx.AsyncClient() as client:
            return await fetch_supplement_zip("PMC1111111", client=client)

    assert asyncio.run(run()) is None


def test_fetch_supplement_zip_returns_none_on_non_404_http_error(httpx_mock: HTTPXMock):
    httpx_mock.add_response(
        url=EUROPEPMC_SUPPLEMENTARY_FILES_URL.format(pmcid="PMC2222222"), status_code=500
    )

    async def run() -> bytes | None:
        async with httpx.AsyncClient() as client:
            return await fetch_supplement_zip("PMC2222222", client=client)

    assert asyncio.run(run()) is None


def test_fetch_supplements_returns_unpacked_files(httpx_mock: HTTPXMock):
    zip_bytes = _build_zip({"tiny.csv": b"a,b\n1,2\n"})
    httpx_mock.add_response(
        url=EUROPEPMC_SUPPLEMENTARY_FILES_URL.format(pmcid="PMC8497572"),
        content=zip_bytes,
        headers={"Content-Type": "application/zip"},
    )

    async def run() -> list[SupplementFile]:
        async with httpx.AsyncClient() as client:
            return await fetch_supplements("PMC8497572", client=client)

    files = asyncio.run(run())
    assert len(files) == 1
    assert files[0].filename == "tiny.csv"
    assert files[0].media_type == "csv"


def test_fetch_supplements_returns_empty_list_on_404(httpx_mock: HTTPXMock):
    httpx_mock.add_response(
        url=EUROPEPMC_SUPPLEMENTARY_FILES_URL.format(pmcid="PMC0000000"), status_code=404
    )

    async def run() -> list[SupplementFile]:
        async with httpx.AsyncClient() as client:
            return await fetch_supplements("PMC0000000", client=client)

    assert asyncio.run(run()) == []


# --- supplement_to_text -------------------------------------------------------------------


def _tiny_xlsx_bytes() -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Taxon", "LDA score"])
    ws.append(["Bacteroides", 3.2])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _tiny_docx_bytes() -> bytes:
    document = docx.Document()
    document.add_paragraph("Uncropped gel images for Figure 2.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Sample"
    table.rows[0].cells[1].text = "Value"
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_supplement_to_text_renders_xlsx_sheet_grid():
    f = SupplementFile(filename="s.xlsx", media_type="xlsx", raw_bytes=_tiny_xlsx_bytes())
    text = supplement_to_text(f)
    assert text is not None
    assert "Sheet1" in text
    assert "Bacteroides" in text
    assert "3.2" in text


def test_supplement_to_text_renders_csv_grid():
    raw = b"Taxon,LDA score\nBacteroides,3.2\n"
    f = SupplementFile(filename="s.csv", media_type="csv", raw_bytes=raw)
    text = supplement_to_text(f)
    assert text == "Taxon\tLDA score\nBacteroides\t3.2"


def test_supplement_to_text_renders_tsv_grid():
    raw = b"Taxon\tLDA score\nBacteroides\t3.2\n"
    f = SupplementFile(filename="s.tsv", media_type="tsv", raw_bytes=raw)
    text = supplement_to_text(f)
    assert text == "Taxon\tLDA score\nBacteroides\t3.2"


def test_supplement_to_text_renders_docx_paragraphs_and_tables():
    f = SupplementFile(filename="s.docx", media_type="docx", raw_bytes=_tiny_docx_bytes())
    text = supplement_to_text(f)
    assert text is not None
    assert "Uncropped gel images for Figure 2." in text
    assert "Sample\tValue" in text


def test_supplement_to_text_returns_none_for_pdf_image_other():
    for media_type in ("pdf", "image", "other"):
        f = SupplementFile(filename=f"s.{media_type}", media_type=media_type, raw_bytes=b"data")
        assert supplement_to_text(f) is None


def test_supplement_to_text_returns_none_for_malformed_xlsx_without_raising():
    f = SupplementFile(filename="bad.xlsx", media_type="xlsx", raw_bytes=b"not an xlsx file")
    assert supplement_to_text(f) is None


def test_supplement_to_text_returns_none_for_malformed_docx_without_raising():
    f = SupplementFile(filename="bad.docx", media_type="docx", raw_bytes=b"not a docx file")
    assert supplement_to_text(f) is None


# --- supplement_to_model_document ---------------------------------------------------------


def test_supplement_to_model_document_pdf_returns_document_block():
    f = SupplementFile(filename="s.pdf", media_type="pdf", raw_bytes=b"%PDF-1.4 fake bytes")
    block = supplement_to_model_document(f)
    assert block is not None
    assert block["type"] == "file"
    assert block["file"]["file_data"].startswith("data:application/pdf;base64,")


def test_supplement_to_model_document_non_pdf_returns_none():
    for media_type in ("xlsx", "csv", "tsv", "docx", "image", "other"):
        f = SupplementFile(filename=f"s.{media_type}", media_type=media_type, raw_bytes=b"data")
        assert supplement_to_model_document(f) is None
