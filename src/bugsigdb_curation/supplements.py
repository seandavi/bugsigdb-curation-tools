"""Supplement identification + fetch + parse (standalone -- not wired into the curator).

For a PMCID, a paper's supplementary files are:

1. **Named** in the EuropePMC ``fullTextXML`` (the same document
   `bugsigdb_curation.retrieval.fetch_fulltext_xml` already fetches), as
   ``<supplementary-material>`` elements, e.g.::

       <supplementary-material id="MOESM1" xmlns:xlink="http://www.w3.org/1999/xlink">
         <label>Supplementary Material 1</label>
         <caption><p>Uncropped gels.</p></caption>
         <media xlink:href="41598_2021_99379_MOESM1_ESM.pdf" mimetype="application" mime-subtype="pdf"/>
       </supplementary-material>

   (:func:`parse_supplement_refs`).
2. **Fetchable as a single ZIP** from EuropePMC's ``supplementaryFiles`` REST
   endpoint (:data:`EUROPEPMC_SUPPLEMENTARY_FILES_URL`,
   :func:`fetch_supplement_zip`), containing every supplementary file (and
   often ancillary figure images) for the article. Not every PMCID has one --
   EuropePMC 404s in that case, tolerated the same way
   `bugsigdb_curation.curator.evidence.assemble_evidence` tolerates a missing
   `fullTextXML` (best-effort, logged, not raised).

This module is deliberately standalone: it is not imported by
`bugsigdb_curation.curator.pipeline` or `bugsigdb_curation.curator.evidence`
(that wiring is a follow-up PR). It also never imports
`bugsigdb_curation.eval` and never reads a gold path, in keeping with the
workflow plan's data firewall (§6e) -- this is a retrieval module, not a
curator module, but there's no reason for it to go anywhere near gold data
either.
"""

from __future__ import annotations

import base64
import csv
import io
import zipfile
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree as ET

import httpx
from loguru import logger

EUROPEPMC_SUPPLEMENTARY_FILES_URL = "https://www.ebi.ac.uk/europepmc/webservices/rest/{pmcid}/supplementaryFiles"

_XLINK_HREF = "{http://www.w3.org/1999/xlink}href"

#: Filename extension -> our small media-type vocabulary. Matched
#: case-insensitively against `Path(filename).suffix`; anything not listed
#: (including no extension) falls back to "image" for common image
#: extensions or "other" otherwise -- see `_media_type_for_filename`.
_EXTENSION_MEDIA_TYPES: dict[str, str] = {
    ".pdf": "pdf",
    ".xlsx": "xlsx",
    ".xls": "xlsx",
    ".csv": "csv",
    ".tsv": "tsv",
    ".docx": "docx",
    ".doc": "docx",
}
_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".tif", ".tiff", ".bmp", ".svg"}


def _media_type_for_filename(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext in _EXTENSION_MEDIA_TYPES:
        return _EXTENSION_MEDIA_TYPES[ext]
    if ext in _IMAGE_EXTENSIONS:
        return "image"
    return "other"


def _text_content(elem: ET.Element) -> str:
    """Flatten an element's text (incl. nested tags) to a string."""
    return "".join(elem.itertext())


def _normalize_ws(text: str) -> str:
    return " ".join(text.split())


@dataclass(frozen=True, slots=True)
class SupplementRef:
    """One `<supplementary-material>` entry named in a fullTextXML document.

    This is metadata only -- the filename it names, plus label/caption --
    not the file's bytes (those come from the separate ZIP fetch, see
    :func:`fetch_supplement_zip`).
    """

    supplement_id: str  # `<supplementary-material id="...">`, or a synthesized "supp-N" fallback
    filename: str | None  # `<media xlink:href="...">`, None if absent
    label: str  # `<label>` text, "" if absent
    caption: str  # `<caption>` text, tags stripped, whitespace collapsed, "" if absent


def parse_supplement_refs(xml_text: str) -> list[SupplementRef]:
    """Parse all `<supplementary-material>` elements out of a fullTextXML document.

    Mirrors `bugsigdb_curation.retrieval.parse_fulltext_figures`'s approach:
    same JATS namespace handling (`xlink:href` via the fully-qualified tag
    name), same "still produce an entry even if a piece is missing" policy.
    Returns `[]` for a document with none.
    """
    root = ET.fromstring(xml_text)
    entries: list[SupplementRef] = []
    for index, supp in enumerate(root.iter("supplementary-material")):
        # `<label>`/`<caption>` are searched as DESCENDANTS, not direct
        # children: real EuropePMC JATS nests the caption inside the
        # `<media>` element (`<supplementary-material><media><caption><p>..`),
        # often with inline markup like `<bold>Additional file 1:</bold>`
        # (which `_text_content`'s itertext flattens). A direct-child
        # `supp.find("caption")` silently returned "" on every real document.
        label_el = supp.find(".//label")
        label = _normalize_ws(_text_content(label_el)) if label_el is not None else ""

        caption_el = supp.find(".//caption")
        caption = _normalize_ws(_text_content(caption_el)) if caption_el is not None else ""

        media_el = supp.find(".//media")
        filename = media_el.get(_XLINK_HREF) if media_el is not None else None

        supplement_id = supp.get("id") or f"supp-{index}"
        entries.append(
            SupplementRef(
                supplement_id=supplement_id,
                filename=filename,
                label=label,
                caption=caption,
            )
        )
    return entries


@dataclass(frozen=True, slots=True)
class SupplementFile:
    """One unpacked supplementary file, with its raw bytes."""

    filename: str
    media_type: str  # "pdf" | "xlsx" | "csv" | "tsv" | "docx" | "image" | "other"
    raw_bytes: bytes


def unpack_supplement_zip(zip_bytes: bytes) -> list[SupplementFile]:
    """Unzip a supplementary-files archive in-memory into `SupplementFile`s.

    Directory entries are skipped. `media_type` is derived from the entry's
    filename extension (see `_media_type_for_filename`) -- the zip's own
    entries carry no separate content-type metadata.
    """
    files: list[SupplementFile] = []
    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            filename = Path(info.filename).name
            raw_bytes = zf.read(info.filename)
            files.append(
                SupplementFile(
                    filename=filename,
                    media_type=_media_type_for_filename(filename),
                    raw_bytes=raw_bytes,
                )
            )
    return files


# --- thin network I/O (not covered by pure-parser tests) --------------------------------


async def fetch_supplement_zip(pmcid: str, *, client: httpx.AsyncClient) -> bytes | None:
    """GET the EuropePMC supplementary-files ZIP for `pmcid`, or None if unavailable.

    Best-effort, mirroring `assemble_evidence`'s fullTextXML 404 tolerance
    (`bugsigdb_curation.curator.evidence`): a 404 (no supplementary files for
    this PMCID) is a normal outcome, not a failure -- logged at INFO and
    returns None. Any other HTTP error, or a 200 response whose
    `Content-Type` isn't a zip, is logged as a WARNING and also returns None
    rather than raising -- this is a best-effort enrichment channel, never
    something that should abort a caller's run.
    """
    log = logger.bind(stage="supplements")
    url = EUROPEPMC_SUPPLEMENTARY_FILES_URL.format(pmcid=pmcid)
    try:
        response = await client.get(url)
        response.raise_for_status()
    except httpx.HTTPStatusError as exc:
        if exc.response.status_code == 404:
            log.info("no supplementary files for pmcid", pmcid=pmcid)
        else:
            log.warning(
                "supplementary files fetch failed",
                pmcid=pmcid,
                status_code=exc.response.status_code,
            )
        return None
    except httpx.HTTPError as exc:
        log.warning("supplementary files fetch failed", pmcid=pmcid, error=str(exc))
        return None

    content_type = response.headers.get("content-type", "")
    if "zip" not in content_type.lower():
        log.warning(
            "supplementary files response was not a zip",
            pmcid=pmcid,
            content_type=content_type,
        )
        return None
    return response.content


async def fetch_supplements(pmcid: str, *, client: httpx.AsyncClient) -> list[SupplementFile]:
    """High-level: fetch the supplementary-files ZIP for `pmcid` and unpack it.

    Returns `[]` (not an error) when there's no ZIP to unpack -- either
    because EuropePMC has none for this PMCID, or the fetch otherwise failed
    best-effort (see `fetch_supplement_zip`).
    """
    zip_bytes = await fetch_supplement_zip(pmcid, client=client)
    if zip_bytes is None:
        return []
    return unpack_supplement_zip(zip_bytes)


# --- model-ready content: text rendering + document content blocks ----------------------


def supplement_to_text(f: SupplementFile) -> str | None:
    """Render a tabular/doc supplementary file as plain text for an LLM's text context.

    - "xlsx": every sheet -> a text grid (sheet-name header line, then
      tab-separated rows).
    - "csv"/"tsv": a single text grid (tab-joined rows).
    - "docx": paragraphs, then table cell text, newline-joined.
    - "pdf"/"image"/"other": None -- PDFs are handled as native document
      blobs (see `supplement_to_model_document`), not text; images have no
      text rendering here at all.

    Defensive: a malformed file logs a warning and returns None rather than
    raising -- this is best-effort enrichment, the same policy as the fetch
    functions above.
    """
    log = logger.bind(stage="supplements")
    try:
        if f.media_type == "xlsx":
            return _xlsx_to_text(f.raw_bytes)
        if f.media_type in ("csv", "tsv"):
            return _delimited_to_text(f.raw_bytes, delimiter="\t" if f.media_type == "tsv" else ",")
        if f.media_type == "docx":
            return _docx_to_text(f.raw_bytes)
    except Exception as exc:  # noqa: BLE001 -- malformed input must degrade to None, never raise
        log.warning("failed to render supplement to text", filename=f.filename, error=str(exc))
        return None
    return None


def _xlsx_to_text(raw_bytes: bytes) -> str:
    import openpyxl

    workbook = openpyxl.load_workbook(io.BytesIO(raw_bytes), read_only=True, data_only=True)
    try:
        sheets_text = []
        for sheet in workbook.worksheets:
            lines = [f"# {sheet.title}"]
            for row in sheet.iter_rows(values_only=True):
                lines.append("\t".join("" if cell is None else str(cell) for cell in row))
            sheets_text.append("\n".join(lines))
        return "\n\n".join(sheets_text)
    finally:
        workbook.close()


def _delimited_to_text(raw_bytes: bytes, *, delimiter: str) -> str:
    text = raw_bytes.decode("utf-8-sig", errors="replace")
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    return "\n".join("\t".join(row) for row in reader)


def _docx_to_text(raw_bytes: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(raw_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(part for part in parts if part is not None)


def supplement_to_model_document(f: SupplementFile) -> dict | None:
    """Render a PDF supplementary file as a LiteLLM document content block.

    Returns `{"type": "file", "file": {"file_data": "data:application/pdf;base64,<...>"}}`
    for "pdf"; None for every other media type (they either have a text
    rendering via `supplement_to_text`, or no model-ready rendering at all,
    e.g. "image"/"other").
    """
    if f.media_type != "pdf":
        return None
    encoded = base64.b64encode(f.raw_bytes).decode("ascii")
    return {"type": "file", "file": {"file_data": f"data:application/pdf;base64,{encoded}"}}
